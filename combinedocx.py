from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK
import sys, os
import glob

def combine_docx_from_folder(folder, output_file="combined.docx"):
    """
    Combine all .docx files from a folder into a single Word document.
    Files are combined in alphabetical order.
    """
    files = sorted(
        f for f in glob.glob(os.path.join(folder, "*.docx"))
        if not os.path.basename(f).startswith("~$")
    )
    
    if not files:
        raise ValueError(f"No .docx files found in {folder}")

    print(f"Found {len(files)} files. Combining into {output_file}...")

    # Start with the first document
    master = Document(files[0])
    composer = Composer(master)

    for idx, file in enumerate(files[1:], start=2):
        try:
            sub_doc = Document(file)
        except Exception as e:
            print(f"⚠️ Skipping {file} due to error: {e}")
            continue
        master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        composer.append(sub_doc)
    composer.save(output_file)
    print(f"✅ Combined document saved as {output_file}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python combine_docx.py <folder_path>")
        sys.exit(1)

    folder_path = sys.argv[1]

    if not os.path.isdir(folder_path):
        print(f"Error: {folder_path} is not a valid folder")
        sys.exit(1)

    combine_docx_from_folder(folder_path, "combined.docx")