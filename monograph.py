import sys
import os
import pandas as pd
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO


def get_unique_folder(base_folder):
    """
    Ensure the output folder is unique.
    If 'monographs' already exists, create 'monographs1', then 'monographs2', etc.
    """
    folder = base_folder
    counter = 1
    # Keep incrementing until a non-existing folder name is found
    while os.path.exists(folder):
        folder = f"{base_folder}{counter}"
        counter += 1
    os.makedirs(folder)  # Create the folder
    return folder


def get_unique_filename(folder, filename):
    """
    Ensure the file name is unique within the output folder.
    If 'file.docx' exists, save as 'file(1).docx', then 'file(2).docx', etc.
    """
    base, ext = os.path.splitext(filename)
    if not ext:  # If no extension is given, default to .docx
        ext = ".docx"
    unique_name = base + ext
    counter = 1
    # Keep checking until a free name is available
    while os.path.exists(os.path.join(folder, unique_name)):
        unique_name = f"{base}({counter}){ext}"
        counter += 1
    return unique_name


def fetch_image(filename, image_folder="images"):
    """
    Fetch an image from a local folder by filename.
    Returns the file path of image, or None if not found.
    """
    filepath = os.path.join(image_folder, filename)
    if os.path.exists(filepath) and os.path.isfile(filepath):
        return filepath
    else:
        print(f"⚠️ Image not found: {filepath}")
        return None


def replace_placeholders_in_table(doc, row_data):
    """
    Replace placeholders inside all tables of the document.
    Placeholders are written like {column_name} in the template.
    
    - For text columns: replace {col} with the value.
    - For 'image' column: fetch the image from Google Drive and insert it.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for col, value in row_data.items():
                    placeholder = f"{{{{{col}}}}}"  # Placeholder format e.g. {{name}}
                    if placeholder in cell.text:
                        if col.lower() == "image":
                            # Replace with image from Drive
                            img_path = fetch_image(str(value).strip())
                            cell.text = ""  # Clear placeholder text
                            if img_path:
                                valid_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp")
                                if img_path.lower().endswith(valid_exts):
                                    try:
                                        print(f"👉 Checking image: {img_path}")
                                        with Image.open(img_path) as im:
                                            im.verify()
                                        with Image.open(img_path) as im:
                                            buf = BytesIO()
                                            if img_path.lower().endswith(".webp"):
                                                print(f"🔄 Converting unsupported format to JPEG: {img_path}")
                                                rgb_im = im.convert("RGB")
                                                rgb_im.save(buf, format="JPEG")
                                            else:
                                                im.save(buf, format=im.format)
                                            buf.seek(0)
                                        run = cell.paragraphs[0].add_run()
                                        run.add_picture(buf, width=Inches(1.5))
                                    except Exception as e:
                                        print(f"❌ Invalid image file skipped: {img_path}")
                                        import traceback
                                        traceback.print_exc()
                                        cell.text = "[Invalid image file]"
                                        continue
                                else:
                                    print(f"❌ Skipping unsupported image format: {img_path}")
                                    cell.text = "[Unsupported image format]"
                            else:
                                cell.text = "[Image fetch failed]"
                        else:
                            # Replace text placeholders
                            cell.text = cell.text.replace(placeholder, str(value))


def fill_template(template_path, csv_path, output_folder="monographs"):
    """
    Main function:
    - Reads the CSV data.
    - Creates a unique output folder (monographs, monographs1, etc.).
    - For each row in the CSV:
        * Load the template
        * Replace placeholders with row values
        * Save the filled document with the filename from the CSV
    """
    # Load CSV into DataFrame
    df = pd.read_csv(csv_path)

    # Create a unique output folder
    output_folder = get_unique_folder(output_folder)

    # Ensure filename column exists in CSV
    if "filename" not in df.columns:
        raise ValueError("CSV must have a column named 'filename' for output file names.")

    # Process each row of the CSV
    for _, row in df.iterrows():
        doc = Document(template_path)
        replace_placeholders(doc, row.to_dict())

        # Generate safe file name
        raw_filename = str(row["filename"]).strip()
        unique_filename = get_unique_filename(output_folder, raw_filename)
        output_file = os.path.join(output_folder, unique_filename)

        # Save the filled document
        doc.save(output_file)
        print(f"✅ Saved: {output_file}")

def replace_placeholders_in_paragraphs(doc, row_data):
    """
    Replace placeholders in normal paragraphs (outside of tables).
    """
    for para in doc.paragraphs:
        for col, value in row_data.items():
            placeholder = f"{{{{{col}}}}}"
            if placeholder in para.text:
                if col.lower() == "image":
                    # Insert image inline instead of text
                    img_path = fetch_image(str(value).strip())
                    para.clear()  # Clear existing text
                    if img_path:
                        valid_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp")
                        if img_path.lower().endswith(valid_exts):
                            try:
                                print(f"👉 Checking image: {img_path}")
                                with Image.open(img_path) as im:
                                    im.verify()
                                with Image.open(img_path) as im:
                                    buf = BytesIO()
                                    if img_path.lower().endswith(".webp"):
                                        print(f"🔄 Converting unsupported format to JPEG: {img_path}")
                                        rgb_im = im.convert("RGB")
                                        rgb_im.save(buf, format="JPEG")
                                    else:
                                        im.save(buf, format=im.format)
                                    buf.seek(0)
                                run = para.add_run()
                                run.add_picture(buf, width=Inches(1.5))
                            except Exception as e:
                                print(f"❌ Invalid image file skipped: {img_path}")
                                import traceback
                                traceback.print_exc()
                                para.add_run("[Invalid image file]")
                                continue
                        else:
                            print(f"❌ Skipping unsupported image format: {img_path}")
                            para.add_run("[Unsupported image format]")
                    else:
                        para.add_run("[Image fetch failed]")
                else:
                    para.text = para.text.replace(placeholder, str(value))

def replace_placeholders(doc, row_data):
    replace_placeholders_in_paragraphs(doc, row_data)
    replace_placeholders_in_table(doc, row_data)

if __name__ == "__main__":
    # Expecting: python script.py template.docx data.csv
    if len(sys.argv) < 3:
        print("Usage: python script.py <template.docx> <data.csv>")
        sys.exit(1)

    template_path = sys.argv[1]
    csv_path = sys.argv[2]

    fill_template(template_path, csv_path, "monographs")